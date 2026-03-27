"""Excel export endpoint — generates comprehensive .xlsx from deal data sent by frontend."""

import io
import math
from datetime import datetime
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel as PydanticBaseModel
from typing import Optional, List, Dict, Any
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.utils import get_column_letter
from openpyxl.workbook.defined_name import DefinedName

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

router = APIRouter(prefix="/api/v1/export", tags=["export"])


# ═══════════════════════════════════════════════════════════
# REQUEST MODEL
# ═══════════════════════════════════════════════════════════

class ExportRequest(PydanticBaseModel):
    """Deal data sent from the frontend for Excel export."""
    property_name: str = "Untitled Deal"
    address: Optional[str] = None
    property_type: Optional[str] = None
    asking_price: Optional[float] = None
    noi: Optional[float] = None
    cap_rate: Optional[float] = None
    rentable_sf: Optional[float] = None
    occupancy_rate: Optional[float] = None
    operating_expenses: Optional[float] = None
    annual_revenue: Optional[float] = None
    year_built: Optional[int] = None
    assumptions: Optional[Dict[str, Any]] = None
    rent_roll: Optional[List[Dict[str, Any]]] = None
    source: Optional[str] = None
    argus_detail: Optional[Dict[str, Any]] = None


# ═══════════════════════════════════════════════════════════
# FINANCIAL HELPERS — mirror frontend JS calculations
# ═══════════════════════════════════════════════════════════

def _monthly_payment(principal: float, annual_rate: float, amort_years: int) -> float:
    r = annual_rate / 12
    n = amort_years * 12
    if r == 0:
        return principal / n if n > 0 else 0
    return principal * (r * (1 + r) ** n) / ((1 + r) ** n - 1)


def _annual_debt_service(principal: float, annual_rate: float, amort_years: int) -> float:
    return _monthly_payment(principal, annual_rate, amort_years) * 12


def _annual_debt_service_io(principal: float, annual_rate: float) -> float:
    return principal * annual_rate


def _loan_balance_at_year(principal: float, annual_rate: float, amort_years: int, year_n: int) -> float:
    r = annual_rate / 12
    n = amort_years * 12
    p = year_n * 12
    pmt = _monthly_payment(principal, annual_rate, amort_years)
    if r == 0:
        return principal - (pmt * p)
    return principal * (1 + r) ** p - pmt * ((1 + r) ** p - 1) / r


def _calculate_irr(cash_flows: list, guess: float = 0.10) -> Optional[float]:
    rate = guess
    for _ in range(200):
        npv = 0.0
        dnpv = 0.0
        for t, cf in enumerate(cash_flows):
            denom = (1 + rate) ** t
            if denom == 0:
                break
            npv += cf / denom
            dnpv -= t * cf / ((1 + rate) ** (t + 1))
        if abs(dnpv) < 1e-12:
            break
        new_rate = rate - npv / dnpv
        if abs(new_rate - rate) < 1e-8:
            return new_rate
        rate = new_rate
    return rate if abs(rate) < 10 else None


def _safe_pct(val, default=0):
    """Normalize a percentage value — if >1 treat as already x100."""
    if val is None:
        return default
    return val / 100 if val > 1 else val


def _fmt_currency(val):
    """Format a number as currency string for text contexts."""
    if val is None:
        return "$0"
    if abs(val) >= 1_000_000:
        return f"${val / 1_000_000:,.1f}M"
    elif abs(val) >= 1_000:
        return f"${val / 1_000:,.0f}K"
    return f"${val:,.0f}"


def _fmt_pct(val):
    """Format decimal as percentage string."""
    if val is None:
        return "0.0%"
    return f"{val * 100:.1f}%" if val < 1 else f"{val:.1f}%"


# ═══════════════════════════════════════════════════════════
# STYLES
# ═══════════════════════════════════════════════════════════

HEADER_FONT = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
HEADER_FILL = PatternFill(start_color="0A1628", end_color="0A1628", fill_type="solid")
SUBHEADER_FILL = PatternFill(start_color="132A42", end_color="132A42", fill_type="solid")
SUBHEADER_FONT = Font(name="Calibri", bold=True, size=11, color="E8EDF2")
LABEL_FONT = Font(name="Calibri", bold=True, size=10, color="8A9BB0")
VALUE_FONT = Font(name="Calibri", size=10)
TITLE_FONT = Font(name="Calibri", bold=True, size=14, color="0A1628")
GREEN_FILL = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
RED_FILL = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
AMBER_FILL = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
GREEN_FONT = Font(name="Calibri", size=10, color="166534")
RED_FONT = Font(name="Calibri", size=10, color="991B1B")
BLUE_FILL = PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid")
METRIC_FILL = PatternFill(start_color="EFF6FF", end_color="EFF6FF", fill_type="solid")
METRIC_FONT = Font(name="Calibri", bold=True, size=10, color="1E3A5F")

CURRENCY_FMT = '#,##0'
CURRENCY_CENTS_FMT = '#,##0.00'
PERCENT_FMT = '0.00%'
NUMBER_FMT = '#,##0'
RATIO_FMT = '0.00"x"'
THIN_BORDER = Border(
    left=Side(style="thin", color="D0D0D0"),
    right=Side(style="thin", color="D0D0D0"),
    top=Side(style="thin", color="D0D0D0"),
    bottom=Side(style="thin", color="D0D0D0"),
)


def _style_header_row(ws, row, num_cols):
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER


def _add_kv_row(ws, row, label, value, value_fmt=None):
    lc = ws.cell(row=row, column=1, value=label)
    lc.font = LABEL_FONT
    lc.border = THIN_BORDER
    vc = ws.cell(row=row, column=2, value=value)
    vc.font = VALUE_FONT
    vc.border = THIN_BORDER
    if value_fmt:
        vc.number_format = value_fmt
    return row + 1


def _section_header(ws, row, title, cols=2):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
    cell = ws.cell(row=row, column=1, value=title)
    cell.font = SUBHEADER_FONT
    cell.fill = SUBHEADER_FILL
    for c in range(1, cols + 1):
        ws.cell(row=row, column=c).fill = SUBHEADER_FILL
        ws.cell(row=row, column=c).border = THIN_BORDER
    return row + 1


# ═══════════════════════════════════════════════════════════
# V1 EXPORT ENDPOINT
# ═══════════════════════════════════════════════════════════

@router.post("")
async def export_deal_to_excel(data: ExportRequest):
    """Generate comprehensive Excel workbook from deal data."""
    wb = Workbook()

    # Parse assumptions
    a = data.assumptions or {}
    ltv = _safe_pct(a.get("ltv"), 0.65)
    interest_rate = _safe_pct(a.get("interest_rate"), 0.055)
    exit_cap = _safe_pct(a.get("exit_cap_rate"), 0.065)
    noi_growth = _safe_pct(a.get("noi_growth"), 0.02)
    hold_period = int(a.get("hold_period", 5))
    amort_years = int(a.get("amortization_years", 25))
    is_io = bool(a.get("interest_only", False))
    io_term = int(a.get("io_term", 5))

    asking_price = data.asking_price or 0
    noi = data.noi or 0
    cap_rate = _safe_pct(data.cap_rate, 0)
    rentable_sf = data.rentable_sf or 0
    debt = asking_price * ltv
    equity = asking_price * (1 - ltv)

    # Debt service
    if debt > 0:
        if is_io:
            ads = _annual_debt_service_io(debt, interest_rate)
        else:
            ads = _annual_debt_service(debt, interest_rate, amort_years)
    else:
        ads = 0

    hold_years = io_term if is_io else hold_period
    dscr = noi / ads if ads > 0 else 0
    coc = (noi - ads) / equity if equity > 0 else 0
    price_per_sf = asking_price / rentable_sf if rentable_sf > 0 else 0
    debt_per_foot = debt / rentable_sf if rentable_sf > 0 else 0
    debt_yield = noi / debt if debt > 0 else 0
    debt_constant = ads / debt if debt > 0 else 0
    leverage_spread = cap_rate - debt_constant

    # ══════════════════════════════════════════
    # Sheet 1: Executive Summary
    # ══════════════════════════════════════════
    ws = wb.active
    ws.title = "Executive Summary"
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 22

    ws.merge_cells("A1:B1")
    ws.cell(row=1, column=1, value=data.property_name).font = TITLE_FONT

    row = 3
    row = _section_header(ws, row, "Property Details")
    if data.address:
        row = _add_kv_row(ws, row, "Address", data.address)
    if data.property_type:
        row = _add_kv_row(ws, row, "Property Type", data.property_type)
    if rentable_sf > 0:
        row = _add_kv_row(ws, row, "Rentable SF", rentable_sf, NUMBER_FMT)
    if data.source:
        row = _add_kv_row(ws, row, "Data Source", data.source.upper())

    row += 1
    row = _section_header(ws, row, "Financial KPIs")
    row = _add_kv_row(ws, row, "Asking Price", asking_price, CURRENCY_FMT)
    row = _add_kv_row(ws, row, "Net Operating Income (NOI)", noi, CURRENCY_FMT)
    row = _add_kv_row(ws, row, "Going-In Cap Rate", cap_rate, PERCENT_FMT)
    if rentable_sf > 0:
        row = _add_kv_row(ws, row, "Price / SF", price_per_sf, CURRENCY_CENTS_FMT)
    row = _add_kv_row(ws, row, "Equity Required", equity, CURRENCY_FMT)
    row = _add_kv_row(ws, row, "Cash-on-Cash (Year 1)", coc, PERCENT_FMT)
    row = _add_kv_row(ws, row, "DSCR", dscr, RATIO_FMT)
    row = _add_kv_row(ws, row, "LTV", ltv, PERCENT_FMT)

    row += 1
    row = _section_header(ws, row, "Debt Metrics")
    row = _add_kv_row(ws, row, "Loan Amount", debt, CURRENCY_FMT)
    if rentable_sf > 0:
        row = _add_kv_row(ws, row, "Debt / Foot", debt_per_foot, CURRENCY_CENTS_FMT)
    row = _add_kv_row(ws, row, "Annual Debt Service", ads, CURRENCY_FMT)
    row = _add_kv_row(ws, row, "Debt Yield", debt_yield, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Debt Constant", debt_constant, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Leverage Spread (Cap - DC)", leverage_spread, PERCENT_FMT)
    if is_io:
        row = _add_kv_row(ws, row, "Interest Only", f"Yes — {io_term} year term")
    else:
        row = _add_kv_row(ws, row, "Interest Only", "No — Amortizing")

    row += 1
    row = _section_header(ws, row, "Underwriting Assumptions")
    row = _add_kv_row(ws, row, "Exit Cap Rate", exit_cap, PERCENT_FMT)
    row = _add_kv_row(ws, row, "NOI Growth Rate", noi_growth, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Hold Period", f"{hold_years} years")
    row = _add_kv_row(ws, row, "Interest Rate", interest_rate, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Amortization", f"{amort_years} years")

    # ══════════════════════════════════════════
    # Sheet 2: Cash Flow Projections
    # ══════════════════════════════════════════
    ws2 = wb.create_sheet("Cash Flow Projections")
    headers = ["Year", "NOI", "Debt Service", "Cap Reserves", "Cash Flow", "Loan Balance", "Cash-on-Cash"]
    col_widths = [8, 16, 16, 16, 16, 18, 14]
    for i, w in enumerate(col_widths):
        ws2.column_dimensions[chr(65 + i)].width = w

    ws2.merge_cells("A1:G1")
    title = f"Cash Flow Projection — {'Interest Only' if is_io else 'Amortizing'}"
    ws2.cell(row=1, column=1, value=title).font = TITLE_FONT

    row = 3
    for ci, h in enumerate(headers, 1):
        cell = ws2.cell(row=row, column=ci, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER
    row += 1

    cap_reserve_psf = 0.40
    total_noi = total_ds = total_cf = 0

    for yr in range(1, hold_years + 1):
        yr_noi = noi * (1 + noi_growth) ** (yr - 1)
        yr_ds = ads
        yr_cap = rentable_sf * cap_reserve_psf
        yr_cf = yr_noi - yr_ds - yr_cap
        if is_io:
            yr_bal = debt
        else:
            yr_bal = _loan_balance_at_year(debt, interest_rate, amort_years, yr) if debt > 0 else 0
        yr_coc = (yr_noi - yr_ds) / equity if equity > 0 else 0

        total_noi += yr_noi
        total_ds += yr_ds
        total_cf += yr_cf

        ws2.cell(row=row, column=1, value=yr).border = THIN_BORDER
        ws2.cell(row=row, column=2, value=yr_noi).number_format = CURRENCY_FMT
        ws2.cell(row=row, column=2).border = THIN_BORDER
        ws2.cell(row=row, column=3, value=yr_ds).number_format = CURRENCY_FMT
        ws2.cell(row=row, column=3).border = THIN_BORDER
        ws2.cell(row=row, column=4, value=yr_cap).number_format = CURRENCY_FMT
        ws2.cell(row=row, column=4).border = THIN_BORDER

        cf_cell = ws2.cell(row=row, column=5, value=yr_cf)
        cf_cell.number_format = CURRENCY_FMT
        cf_cell.border = THIN_BORDER
        cf_cell.font = Font(name="Calibri", bold=True, color="166534" if yr_cf > 0 else "991B1B")

        ws2.cell(row=row, column=6, value=yr_bal).number_format = CURRENCY_FMT
        ws2.cell(row=row, column=6).border = THIN_BORDER

        coc_cell = ws2.cell(row=row, column=7, value=yr_coc)
        coc_cell.number_format = PERCENT_FMT
        coc_cell.border = THIN_BORDER
        if yr_coc > 0.08:
            coc_cell.fill = GREEN_FILL
        elif yr_coc < 0:
            coc_cell.fill = RED_FILL

        row += 1

    # Totals
    for ci in range(1, 8):
        ws2.cell(row=row, column=ci).border = THIN_BORDER
        ws2.cell(row=row, column=ci).font = Font(name="Calibri", bold=True, size=10)
    ws2.cell(row=row, column=1, value="TOTAL")
    ws2.cell(row=row, column=2, value=total_noi).number_format = CURRENCY_FMT
    ws2.cell(row=row, column=3, value=total_ds).number_format = CURRENCY_FMT
    ws2.cell(row=row, column=4, value=rentable_sf * cap_reserve_psf * hold_years).number_format = CURRENCY_FMT
    ws2.cell(row=row, column=5, value=total_cf).number_format = CURRENCY_FMT

    # ══════════════════════════════════════════
    # Sheet 3: IRR Sensitivity Matrix
    # ══════════════════════════════════════════
    ws3 = wb.create_sheet("IRR Sensitivity")
    exit_caps = [0.055, 0.060, 0.065, 0.070, 0.075, 0.080]
    hold_periods = [3, 5, 7, 10]

    ws3.merge_cells("A1:E1")
    ws3.cell(row=1, column=1, value="Levered IRR Sensitivity — Exit Cap x Hold Period").font = TITLE_FONT
    ws3.column_dimensions["A"].width = 16

    row = 3
    # Header
    ws3.cell(row=row, column=1, value="Exit Cap / Hold").font = HEADER_FONT
    ws3.cell(row=row, column=1).fill = HEADER_FILL
    ws3.cell(row=row, column=1).border = THIN_BORDER
    for ci, hp in enumerate(hold_periods, 2):
        cell = ws3.cell(row=row, column=ci, value=f"{hp} Years")
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER
        ws3.column_dimensions[chr(64 + ci)].width = 14
    row += 1

    for ec in exit_caps:
        ws3.cell(row=row, column=1, value=ec).number_format = PERCENT_FMT
        ws3.cell(row=row, column=1).font = Font(name="Calibri", bold=True, size=10)
        ws3.cell(row=row, column=1).border = THIN_BORDER

        for ci, hp in enumerate(hold_periods, 2):
            # Build IRR cash flows
            flows = [-equity] if equity > 0 else [0]
            for yr in range(1, hp + 1):
                yr_noi = noi * (1 + noi_growth) ** (yr - 1)
                yr_ds = ads
                yr_cap = rentable_sf * cap_reserve_psf
                yr_cf = yr_noi - yr_ds - yr_cap
                flows.append(yr_cf)

            # Exit proceeds
            exit_noi = noi * (1 + noi_growth) ** hp
            sale_price = exit_noi / ec if ec > 0 else 0
            if is_io:
                exit_balance = debt
            else:
                exit_balance = _loan_balance_at_year(debt, interest_rate, amort_years, hp) if debt > 0 else 0
            sale_proceeds = sale_price - exit_balance
            flows[-1] += sale_proceeds

            irr = _calculate_irr(flows)
            cell = ws3.cell(row=row, column=ci)
            cell.border = THIN_BORDER
            cell.alignment = Alignment(horizontal="center")
            if irr is not None and abs(irr) < 5:
                cell.value = irr
                cell.number_format = PERCENT_FMT
                if irr > 0.15:
                    cell.fill = GREEN_FILL
                    cell.font = GREEN_FONT
                elif irr < 0:
                    cell.fill = RED_FILL
                    cell.font = RED_FONT
            else:
                cell.value = "N/A"

        row += 1

    # ══════════════════════════════════════════
    # Sheet 4: Property Value Sensitivity
    # ══════════════════════════════════════════
    ws4 = wb.create_sheet("Value Sensitivity")
    growth_rates = [-0.05, -0.025, 0, 0.025, 0.05, 0.075, 0.10]
    exit_cap_rates = [0.050, 0.055, 0.060, 0.065, 0.070, 0.075, 0.080]

    ws4.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(exit_cap_rates) + 1)
    ws4.cell(row=1, column=1, value=f"Property Value at Exit (Year {hold_years}) — NOI Growth x Exit Cap").font = TITLE_FONT
    ws4.column_dimensions["A"].width = 18

    row = 3
    ws4.cell(row=row, column=1, value="Growth / Cap").font = HEADER_FONT
    ws4.cell(row=row, column=1).fill = HEADER_FILL
    ws4.cell(row=row, column=1).border = THIN_BORDER
    for ci, ecr in enumerate(exit_cap_rates, 2):
        cell = ws4.cell(row=row, column=ci, value=ecr)
        cell.number_format = PERCENT_FMT
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER
        ws4.column_dimensions[chr(64 + ci)].width = 16
    row += 1

    for gr in growth_rates:
        ws4.cell(row=row, column=1, value=gr).number_format = PERCENT_FMT
        ws4.cell(row=row, column=1).font = Font(name="Calibri", bold=True, size=10)
        ws4.cell(row=row, column=1).border = THIN_BORDER

        for ci, ecr in enumerate(exit_cap_rates, 2):
            exit_noi = noi * (1 + gr) ** hold_years
            exit_val = exit_noi / ecr if ecr > 0 else 0
            cell = ws4.cell(row=row, column=ci, value=exit_val)
            cell.number_format = CURRENCY_FMT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(horizontal="center")
            if exit_val > asking_price and asking_price > 0:
                cell.fill = GREEN_FILL
                cell.font = GREEN_FONT
            elif asking_price > 0:
                cell.fill = RED_FILL
                cell.font = RED_FONT
        row += 1

    # ══════════════════════════════════════════
    # Sheet 5: Debt Analysis
    # ══════════════════════════════════════════
    ws5 = wb.create_sheet("Debt Analysis")
    ws5.column_dimensions["A"].width = 22
    ws5.column_dimensions["B"].width = 16
    ws5.column_dimensions["C"].width = 16
    ws5.column_dimensions["D"].width = 16
    ws5.column_dimensions["E"].width = 12

    ws5.merge_cells("A1:E1")
    ws5.cell(row=1, column=1, value="Debt Analysis & Leverage Assessment").font = TITLE_FONT

    row = 3
    row = _section_header(ws5, row, "Leverage Analysis", 2)
    row = _add_kv_row(ws5, row, "Going-In Cap Rate", cap_rate, PERCENT_FMT)
    row = _add_kv_row(ws5, row, "Debt Constant", debt_constant, PERCENT_FMT)
    row = _add_kv_row(ws5, row, "Leverage Spread", leverage_spread, PERCENT_FMT)
    if leverage_spread > 0.005:
        row = _add_kv_row(ws5, row, "Leverage Status", "Positive — Favorable")
    elif leverage_spread > -0.005:
        row = _add_kv_row(ws5, row, "Leverage Status", "Neutral")
    else:
        row = _add_kv_row(ws5, row, "Leverage Status", "Negative — Unfavorable")

    # NOI vs Debt Service Table
    row += 1
    row = _section_header(ws5, row, "NOI vs Debt Service Projection", 5)
    headers5 = ["Year", "NOI", "Debt Service", "Free Cash Flow", "DSCR"]
    for ci, h in enumerate(headers5, 1):
        cell = ws5.cell(row=row, column=ci, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER
    row += 1

    t_noi = t_fcf = 0
    for yr in range(1, hold_years + 1):
        yr_noi = noi * (1 + noi_growth) ** (yr - 1)
        fcf = yr_noi - ads
        yr_dscr = yr_noi / ads if ads > 0 else 0
        t_noi += yr_noi
        t_fcf += fcf

        ws5.cell(row=row, column=1, value=yr).border = THIN_BORDER
        ws5.cell(row=row, column=2, value=yr_noi).number_format = CURRENCY_FMT
        ws5.cell(row=row, column=2).border = THIN_BORDER
        ws5.cell(row=row, column=3, value=ads).number_format = CURRENCY_FMT
        ws5.cell(row=row, column=3).border = THIN_BORDER
        ws5.cell(row=row, column=4, value=fcf).number_format = CURRENCY_FMT
        ws5.cell(row=row, column=4).border = THIN_BORDER
        ws5.cell(row=row, column=4).font = Font(name="Calibri", color="166534" if fcf > 0 else "991B1B")
        dscr_cell = ws5.cell(row=row, column=5, value=yr_dscr)
        dscr_cell.number_format = RATIO_FMT
        dscr_cell.border = THIN_BORDER
        if yr_dscr >= 1.25:
            dscr_cell.font = GREEN_FONT
        elif yr_dscr < 1.0:
            dscr_cell.font = RED_FONT
        row += 1

    # Totals
    for ci in range(1, 6):
        ws5.cell(row=row, column=ci).border = THIN_BORDER
        ws5.cell(row=row, column=ci).font = Font(name="Calibri", bold=True, size=10)
    ws5.cell(row=row, column=1, value="TOTAL")
    ws5.cell(row=row, column=2, value=t_noi).number_format = CURRENCY_FMT
    ws5.cell(row=row, column=3, value=ads * hold_years).number_format = CURRENCY_FMT
    ws5.cell(row=row, column=4, value=t_fcf).number_format = CURRENCY_FMT

    # Debt Constant Reference Grid
    row += 2
    ref_rates = [3.0, 4.0, 5.0, 5.5, 6.0, 6.5, 7.0, 7.5, 8.0]
    ref_amorts = [15, 20, 25, 30]

    row = _section_header(ws5, row, "Debt Constant Reference (Rate x Amortization)", 1 + len(ref_amorts))
    ws5.cell(row=row, column=1, value="Rate / Amort").font = HEADER_FONT
    ws5.cell(row=row, column=1).fill = HEADER_FILL
    ws5.cell(row=row, column=1).border = THIN_BORDER
    for ci, am in enumerate(ref_amorts, 2):
        cell = ws5.cell(row=row, column=ci, value=f"{am} yr")
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER
    row += 1

    for rate in ref_rates:
        ws5.cell(row=row, column=1, value=f"{rate:.1f}%").border = THIN_BORDER
        ws5.cell(row=row, column=1).font = Font(name="Calibri", bold=True, size=10)
        for ci, am in enumerate(ref_amorts, 2):
            if rate > 0:
                dc = _annual_debt_service(1.0, rate / 100, am)
            else:
                dc = 0
            cell = ws5.cell(row=row, column=ci, value=dc)
            cell.number_format = PERCENT_FMT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(horizontal="center")
            # Highlight current selection
            if abs(rate - interest_rate * 100) < 0.2 and am == amort_years:
                cell.fill = PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid")
                cell.font = Font(name="Calibri", bold=True, size=10, color="1E40AF")
        row += 1

    # ══════════════════════════════════════════
    # Sheet 6: Rent Roll
    # ══════════════════════════════════════════
    if data.rent_roll:
        ws6 = wb.create_sheet("Rent Roll")
        rr_headers = ["Unit", "Tenant", "SF", "Rent PSF", "Annual Rent", "Monthly Rent", "Lease End"]
        rr_keys = ["unit", "tenant", "sf", "rent_psf", "annual_rent", "monthly_rent", "expiry"]
        col_widths6 = [14, 30, 12, 12, 16, 16, 14]

        for i, w in enumerate(col_widths6):
            ws6.column_dimensions[chr(65 + i)].width = w

        for ci, h in enumerate(rr_headers, 1):
            cell = ws6.cell(row=1, column=ci, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER

        for ri, entry in enumerate(data.rent_roll, 2):
            for ci, key in enumerate(rr_keys, 1):
                val = entry.get(key, "")
                # Fallback for lease_end / expiry naming
                if key == "expiry" and not val:
                    val = entry.get("lease_end", "")
                if key == "monthly_rent" and not val:
                    annual = entry.get("annual_rent", 0) or 0
                    val = annual / 12 if annual > 0 else 0
                cell = ws6.cell(row=ri, column=ci, value=val)
                cell.border = THIN_BORDER
                if key == "sf":
                    cell.number_format = NUMBER_FMT
                elif key == "rent_psf":
                    cell.number_format = CURRENCY_CENTS_FMT
                elif key in ("annual_rent", "monthly_rent"):
                    cell.number_format = CURRENCY_FMT

        # Summary row
        total_row = len(data.rent_roll) + 2
        ws6.cell(row=total_row, column=1, value="TOTAL").font = Font(name="Calibri", bold=True, size=10)
        ws6.cell(row=total_row, column=1).border = THIN_BORDER

        total_sf = sum(e.get("sf", 0) or 0 for e in data.rent_roll)
        total_annual = sum(e.get("annual_rent", 0) or 0 for e in data.rent_roll)
        total_monthly = sum((e.get("monthly_rent") or (e.get("annual_rent", 0) or 0) / 12) for e in data.rent_roll)
        wa_rent = total_annual / total_sf if total_sf > 0 else 0

        ws6.cell(row=total_row, column=3, value=total_sf).number_format = NUMBER_FMT
        ws6.cell(row=total_row, column=3).font = Font(name="Calibri", bold=True)
        ws6.cell(row=total_row, column=3).border = THIN_BORDER
        ws6.cell(row=total_row, column=4, value=wa_rent).number_format = CURRENCY_CENTS_FMT
        ws6.cell(row=total_row, column=4).font = Font(name="Calibri", bold=True)
        ws6.cell(row=total_row, column=4).border = THIN_BORDER
        ws6.cell(row=total_row, column=5, value=total_annual).number_format = CURRENCY_FMT
        ws6.cell(row=total_row, column=5).font = Font(name="Calibri", bold=True)
        ws6.cell(row=total_row, column=5).border = THIN_BORDER
        ws6.cell(row=total_row, column=6, value=total_monthly).number_format = CURRENCY_FMT
        ws6.cell(row=total_row, column=6).font = Font(name="Calibri", bold=True)
        ws6.cell(row=total_row, column=6).border = THIN_BORDER

        # Occupancy summary
        occupied_sf = sum(e.get("sf", 0) or 0 for e in data.rent_roll if (e.get("tenant", "") or "").lower() != "vacant" and (e.get("annual_rent", 0) or 0) > 0)
        occ_pct = occupied_sf / total_sf if total_sf > 0 else 0

        sr = total_row + 2
        ws6.cell(row=sr, column=1, value="Occupied SF").font = LABEL_FONT
        ws6.cell(row=sr, column=2, value=occupied_sf).number_format = NUMBER_FMT
        ws6.cell(row=sr + 1, column=1, value="Vacant SF").font = LABEL_FONT
        ws6.cell(row=sr + 1, column=2, value=total_sf - occupied_sf).number_format = NUMBER_FMT
        ws6.cell(row=sr + 2, column=1, value="Occupancy Rate").font = LABEL_FONT
        ws6.cell(row=sr + 2, column=2, value=occ_pct).number_format = PERCENT_FMT
        ws6.cell(row=sr + 3, column=1, value="Weighted Avg Rent/SF").font = LABEL_FONT
        ws6.cell(row=sr + 3, column=2, value=wa_rent).number_format = CURRENCY_CENTS_FMT

    # ── Write to buffer ──
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    # Sanitize filename for HTTP headers (latin-1 safe)
    safe_name = data.property_name.replace(" ", "_").replace("/", "-")
    safe_name = safe_name.encode("ascii", "ignore").decode("ascii")[:40]
    filename = f"{safe_name}_Analysis.xlsx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════════════════════
# V2 EXPORT — Enhanced DCF with Waterfall & Value-Add
# ═══════════════════════════════════════════════════════════

class V2ExportRequest(PydanticBaseModel):
    """V2 deal data including computed model results."""
    v2_state: Optional[Dict[str, Any]] = None
    calc: Optional[Dict[str, Any]] = None


def _v2_build_tenant_analytics(tenants: list, total_sf: float):
    """Build tenant analytics data used by multiple sheets."""
    if not tenants:
        return {}

    # Parse expiry years
    expiry_buckets = {}  # year -> list of tenants
    for t in tenants:
        end_str = t.get("end", "")
        yr = None
        if end_str:
            try:
                if len(end_str) >= 4:
                    yr = int(end_str[:4])
            except (ValueError, TypeError):
                pass
        if yr:
            expiry_buckets.setdefault(yr, []).append(t)

    # Top tenants by revenue
    tenants_with_rev = []
    for t in tenants:
        sf = t.get("sf", 0) or 0
        rent_psf = t.get("rentPSF", 0) or 0
        annual_rev = sf * rent_psf
        tenants_with_rev.append({**t, "_annual_rev": annual_rev, "_sf": sf})
    tenants_with_rev.sort(key=lambda x: x["_annual_rev"], reverse=True)

    total_revenue = sum(t["_annual_rev"] for t in tenants_with_rev)
    top_tenant_pct = tenants_with_rev[0]["_annual_rev"] / total_revenue if total_revenue > 0 and tenants_with_rev else 0

    # WALT calculation
    now_year = datetime.now().year
    weighted_years = 0
    total_weight_sf = 0
    for t in tenants:
        sf = t.get("sf", 0) or 0
        end_str = t.get("end", "")
        if end_str and sf > 0:
            try:
                end_yr = int(end_str[:4])
                remaining = max(0, end_yr - now_year)
                weighted_years += remaining * sf
                total_weight_sf += sf
            except (ValueError, TypeError):
                pass
    walt = weighted_years / total_weight_sf if total_weight_sf > 0 else 0

    return {
        "expiry_buckets": expiry_buckets,
        "tenants_with_rev": tenants_with_rev,
        "total_revenue": total_revenue,
        "top_tenant_pct": top_tenant_pct,
        "walt": walt,
    }


@router.post("/v2")
async def export_v2_deal_to_excel(data: V2ExportRequest):
    """Generate institutional-grade V2 Excel workbook with formulas, charts and analytics."""

    wb = Workbook()

    state = data.v2_state or {}
    calc = data.calc or {}
    p = state.get("assumptions", {})
    wf = state.get("waterfall", {})
    tenants = state.get("tenants", [])
    events = state.get("valueAddEvents", [])
    cap_items = state.get("capexItems", [])
    years = calc.get("years", [])

    prop_name = p.get("name", "Untitled Deal")
    hold = p.get("holdPeriod", 5)
    total_sf = p.get("sf", 0) or 0

    # Pre-compute tenant analytics for reuse
    tenant_analytics = _v2_build_tenant_analytics(tenants, total_sf)

    # ══════════════════════════════════════════
    # Sheet 1: Assumptions (with named ranges)
    # ══════════════════════════════════════════
    ws = wb.active
    ws.title = "Assumptions"
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 20

    ws.merge_cells("A1:B1")
    ws.cell(row=1, column=1, value=prop_name).font = TITLE_FONT

    row = 3

    # Property Section
    row = _section_header(ws, row, "Property", 2)
    row = _add_kv_row(ws, row, "Name", p.get("name", ""))
    row = _add_kv_row(ws, row, "Type", p.get("assetType", ""))
    row = _add_kv_row(ws, row, "Address", p.get("address", ""))
    sf_row = row
    row = _add_kv_row(ws, row, "Rentable SF", p.get("sf", 0), NUMBER_FMT)
    # Create named range for SF
    sf_cell = f"B{sf_row}"
    try:
        wb.defined_names.add(DefinedName("SF", attr_text=f"Assumptions!{sf_cell}"))
    except Exception:
        pass

    row += 1

    # Acquisition Section
    row = _section_header(ws, row, "Acquisition", 2)
    pp_row = row
    purchase_price = p.get("purchasePrice", calc.get("pp", 0))
    row = _add_kv_row(ws, row, "Purchase Price", purchase_price, CURRENCY_FMT)
    try:
        wb.defined_names.add(DefinedName("PurchasePrice", attr_text=f"Assumptions!$B${pp_row}"))
    except Exception:
        pass

    row = _add_kv_row(ws, row, "Acq Cost %", (p.get("acqCostPct", 0) or 0) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Closing Costs", (p.get("closingCosts", 0) or 0), CURRENCY_FMT)

    row += 1

    # Financing Section
    row = _section_header(ws, row, "Financing", 2)
    ltv_row = row
    ltv_pct = (p.get("ltv", 65) or 65)
    row = _add_kv_row(ws, row, "LTV", ltv_pct / 100, PERCENT_FMT)
    try:
        wb.defined_names.add(DefinedName("LTV", attr_text=f"Assumptions!$B${ltv_row}"))
    except Exception:
        pass

    rate_row = row
    row = _add_kv_row(ws, row, "Interest Rate", (p.get("rate", 6) or 6) / 100, PERCENT_FMT)
    try:
        wb.defined_names.add(DefinedName("Rate", attr_text=f"Assumptions!$B${rate_row}"))
    except Exception:
        pass

    amort_row = row
    row = _add_kv_row(ws, row, "Amortization (yrs)", p.get("amortYears", 25), NUMBER_FMT)
    try:
        wb.defined_names.add(DefinedName("AmortYears", attr_text=f"Assumptions!$B${amort_row}"))
    except Exception:
        pass

    row = _add_kv_row(ws, row, "Interest Only Period", p.get("ioPeriod", 0), NUMBER_FMT)
    row = _add_kv_row(ws, row, "Origination Fee %", (p.get("origFee", 0) or 0) / 100, PERCENT_FMT)

    # Compute loan amount & equity for named ranges
    loan_amount_val = (purchase_price or 0) * (ltv_pct / 100)
    equity_val = calc.get("totalEq", 0) or calc.get("equity", 0) or ((purchase_price or 0) - loan_amount_val)

    # Add Loan Amount and Equity as named range cells
    row += 1
    row = _section_header(ws, row, "Derived Capital Structure", 2)
    la_row = row
    row = _add_kv_row(ws, row, "Loan Amount", loan_amount_val, CURRENCY_FMT)
    try:
        wb.defined_names.add(DefinedName("LoanAmount", attr_text=f"Assumptions!$B${la_row}"))
    except Exception:
        pass

    eq_row = row
    row = _add_kv_row(ws, row, "Total Equity", equity_val, CURRENCY_FMT)
    try:
        wb.defined_names.add(DefinedName("Equity", attr_text=f"Assumptions!$B${eq_row}"))
    except Exception:
        pass

    row += 1

    # Operations Section
    row = _section_header(ws, row, "Operations", 2)
    noi_assum_row = row
    row = _add_kv_row(ws, row, "Year 1 NOI", p.get("y1NOI", calc.get("y1NOI", 0)), CURRENCY_FMT)
    try:
        wb.defined_names.add(DefinedName("Y1NOI", attr_text=f"Assumptions!$B${noi_assum_row}"))
    except Exception:
        pass

    growth_row = row
    row = _add_kv_row(ws, row, "Rent Growth %", (p.get("rentGrowth", 2) or 2) / 100, PERCENT_FMT)
    try:
        wb.defined_names.add(DefinedName("RentGrowth", attr_text=f"Assumptions!$B${growth_row}"))
    except Exception:
        pass

    row = _add_kv_row(ws, row, "Vacancy %", (p.get("vacancy", 5) or 5) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "CapEx Reserve %", (p.get("capexResPct", 1) or 1) / 100, PERCENT_FMT)

    row += 1

    # Exit Section
    row = _section_header(ws, row, "Exit", 2)
    hold_row = row
    row = _add_kv_row(ws, row, "Hold Period (yrs)", p.get("holdPeriod", 5), NUMBER_FMT)
    try:
        wb.defined_names.add(DefinedName("HoldPeriod", attr_text=f"Assumptions!$B${hold_row}"))
    except Exception:
        pass

    exit_cap_row = row
    row = _add_kv_row(ws, row, "Exit Cap Rate", (p.get("exitCap", 6.5) or 6.5) / 100, PERCENT_FMT)
    try:
        wb.defined_names.add(DefinedName("ExitCap", attr_text=f"Assumptions!$B${exit_cap_row}"))
    except Exception:
        pass

    row = _add_kv_row(ws, row, "Prepay Penalty %", (p.get("prepayPct", 1) or 1) / 100, PERCENT_FMT)

    row += 1

    # Market Leasing Section
    row = _section_header(ws, row, "Market Leasing", 2)
    row = _add_kv_row(ws, row, "Renewal Probability %", (p.get("marketRenewalProb", 75) or 75) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Months Vacant", p.get("marketVacantMonths", 6), NUMBER_FMT)
    row = _add_kv_row(ws, row, "Free Rent Months (New)", p.get("marketFreeRentMonths", 2), NUMBER_FMT)
    row = _add_kv_row(ws, row, "TI New Lease /SF", p.get("marketTINewPSF", 15), CURRENCY_CENTS_FMT)
    row = _add_kv_row(ws, row, "TI Renewal /SF", p.get("marketTIRenewalPSF", 10), CURRENCY_CENTS_FMT)
    row = _add_kv_row(ws, row, "LC New Lease %", (p.get("marketLCNewPct", 6) or 6) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "LC Renewal %", (p.get("marketLCRenewalPct", 4) or 4) / 100, PERCENT_FMT)

    row += 1

    # Hurdles Section
    row = _section_header(ws, row, "Return Hurdles", 2)
    row = _add_kv_row(ws, row, "Tier 1 IRR Threshold", (wf.get("tier1Thresh", 15) or 15) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Tier 2 IRR Threshold", (wf.get("tier2Thresh", 20) or 20) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Tier 2 LP Split %", (wf.get("tier2Split", 70) or 70) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Tier 3 LP Split %", (wf.get("tier3Split", 60) or 60) / 100, PERCENT_FMT)
    row = _add_kv_row(ws, row, "Min EM Threshold", wf.get("minEM", 1.5), RATIO_FMT)

    # ══════════════════════════════════════════
    # Sheet 2: Cash Flows (with formulas)
    # ══════════════════════════════════════════
    ws2 = wb.create_sheet("Cash Flows")
    ws2.column_dimensions["A"].width = 28

    # Set year column widths
    num_years = len(years)
    total_col = num_years + 2  # column index for "Total" column
    for i in range(num_years + 1):
        ws2.column_dimensions[get_column_letter(i + 2)].width = 14

    ws2.merge_cells(f"A1:{get_column_letter(total_col)}1")
    ws2.cell(row=1, column=1, value="Cash Flow Projection (Annual P&L)").font = TITLE_FONT

    # Year headers
    row = 3
    ws2.cell(row=row, column=1, value="Line Item").font = HEADER_FONT
    ws2.cell(row=row, column=1).fill = HEADER_FILL
    ws2.cell(row=row, column=1).border = THIN_BORDER

    for i, y in enumerate(years):
        yr_num = y.get("yr", i + 1)
        col = i + 2
        cell = ws2.cell(row=row, column=col, value=f"Year {yr_num}")
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER

    # Total column header
    cell = ws2.cell(row=row, column=total_col, value="Total")
    cell.font = HEADER_FONT
    cell.fill = HEADER_FILL
    cell.alignment = Alignment(horizontal="center")
    cell.border = THIN_BORDER
    row += 1

    # ── Helper: add a data row with optional SUM total column ──
    def _add_cf_data_row(ws, row_num, label, values, negate=False, fmt=CURRENCY_FMT, add_total=True):
        """Write a cash flow row with values and optional SUM total."""
        ws.cell(row=row_num, column=1, value=label).font = LABEL_FONT
        ws.cell(row=row_num, column=1).border = THIN_BORDER
        for i, val in enumerate(values):
            col = i + 2
            v = -(val or 0) if negate else (val or 0)
            cell = ws.cell(row=row_num, column=col, value=v)
            cell.number_format = fmt
            cell.border = THIN_BORDER
        if add_total and len(values) > 0:
            first_col = get_column_letter(2)
            last_col = get_column_letter(num_years + 1)
            total_cell = ws.cell(row=row_num, column=total_col,
                                 value=f"=SUM({first_col}{row_num}:{last_col}{row_num})")
            total_cell.number_format = fmt
            total_cell.border = THIN_BORDER
        return row_num

    # Revenue section
    row = _section_header(ws2, row, "REVENUE", total_col)

    # Potential Base Rent
    br_row = row
    _add_cf_data_row(ws2, row, "Potential Base Rent", [y.get("baseRent", 0) for y in years])
    row += 1

    # CAM / Expense Recovery
    cam_row = row
    _add_cf_data_row(ws2, row, "CAM / Expense Recovery", [y.get("cam", 0) for y in years])
    row += 1

    # Potential Gross Revenue (formula)
    pgr_row = row
    ws2.cell(row=row, column=1, value="Potential Gross Revenue").font = LABEL_FONT
    ws2.cell(row=row, column=1).border = THIN_BORDER
    for i in range(num_years):
        col = i + 2
        col_letter = get_column_letter(col)
        cell = ws2.cell(row=row, column=col, value=f"={col_letter}{br_row}+{col_letter}{cam_row}")
        cell.number_format = CURRENCY_FMT
        cell.border = THIN_BORDER
    # Total for PGR
    first_col = get_column_letter(2)
    last_col = get_column_letter(num_years + 1)
    ws2.cell(row=row, column=total_col,
             value=f"=SUM({first_col}{pgr_row}:{last_col}{pgr_row})").number_format = CURRENCY_FMT
    ws2.cell(row=row, column=total_col).border = THIN_BORDER
    row += 1

    # Vacancy Loss
    vac_loss_row = row
    _add_cf_data_row(ws2, row, "Vacancy Loss", [y.get("vacancyLoss", 0) for y in years], negate=True)
    row += 1

    # Free Rent
    fr_row = row
    _add_cf_data_row(ws2, row, "Free Rent",
                     [y.get("freeRentLoss", 0) or y.get("freeRent", 0) or 0 for y in years], negate=True)
    row += 1

    # Effective Gross Revenue (formula)
    egr_row = row
    ws2.cell(row=row, column=1, value="Effective Gross Revenue").font = LABEL_FONT
    ws2.cell(row=row, column=1).border = THIN_BORDER
    for i in range(num_years):
        col = i + 2
        cl = get_column_letter(col)
        cell = ws2.cell(row=row, column=col,
                        value=f"={cl}{pgr_row}+{cl}{vac_loss_row}+{cl}{fr_row}")
        cell.number_format = CURRENCY_FMT
        cell.border = THIN_BORDER
    ws2.cell(row=row, column=total_col,
             value=f"=SUM({first_col}{egr_row}:{last_col}{egr_row})").number_format = CURRENCY_FMT
    ws2.cell(row=row, column=total_col).border = THIN_BORDER
    row += 1

    row += 1
    row = _section_header(ws2, row, "NET OPERATING INCOME", total_col)

    # NOI
    noi_row = row
    _add_cf_data_row(ws2, row, "NOI", [y.get("noi", 0) for y in years])
    row += 1

    row += 1
    row = _section_header(ws2, row, "DEBT & CASH FLOW", total_col)

    # Debt Service
    ds_row = row
    _add_cf_data_row(ws2, row, "Debt Service", [y.get("annDS", 0) for y in years], negate=True)
    row += 1

    # CapEx & Reserves
    capex_row = row
    capex_vals = [(y.get("capexRes", 0) or 0) + (y.get("specCapex", 0) or 0) for y in years]
    _add_cf_data_row(ws2, row, "CapEx & Reserves", capex_vals, negate=True)
    row += 1

    # TI / LC
    tilc_row = row
    _add_cf_data_row(ws2, row, "TI / LC", [y.get("tiLC", 0) for y in years], negate=True)
    row += 1

    # Value-Add Income
    vai_row = row
    _add_cf_data_row(ws2, row, "Value-Add Income", [y.get("vaInc", 0) or 0 for y in years])
    row += 1

    # Value-Add Cost
    vac_cost_row = row
    _add_cf_data_row(ws2, row, "Value-Add Cost", [y.get("vaCost", 0) or 0 for y in years], negate=True)
    row += 1

    row += 1
    # Net Cash Flow (formula)
    ncf_row = row
    ws2.cell(row=row, column=1, value="Net Cash Flow").font = Font(name="Calibri", bold=True, size=11)
    ws2.cell(row=row, column=1).border = THIN_BORDER
    ws2.cell(row=row, column=1).fill = AMBER_FILL
    for i in range(num_years):
        col = i + 2
        cl = get_column_letter(col)
        cell = ws2.cell(row=row, column=col,
                        value=f"={cl}{noi_row}+{cl}{ds_row}+{cl}{capex_row}+{cl}{tilc_row}+{cl}{vai_row}+{cl}{vac_cost_row}")
        cell.number_format = CURRENCY_FMT
        cell.border = THIN_BORDER
        cell.fill = AMBER_FILL
        cell.font = Font(name="Calibri", bold=True)

    # NCF total column (SUM formula)
    ncf_total_cell = ws2.cell(row=row, column=total_col,
                               value=f"=SUM({first_col}{ncf_row}:{last_col}{ncf_row})")
    ncf_total_cell.number_format = CURRENCY_FMT
    ncf_total_cell.border = THIN_BORDER
    ncf_total_cell.fill = AMBER_FILL
    ncf_total_cell.font = Font(name="Calibri", bold=True)

    # ── CASH FLOW METRICS section ──
    row += 2
    row = _section_header(ws2, row, "CASH FLOW METRICS", total_col)

    # Cash-on-Cash % row (formula = NCF / Equity)
    coc_row = row
    ws2.cell(row=row, column=1, value="Cash-on-Cash %").font = METRIC_FONT
    ws2.cell(row=row, column=1).border = THIN_BORDER
    ws2.cell(row=row, column=1).fill = METRIC_FILL
    for i in range(num_years):
        col = i + 2
        cl = get_column_letter(col)
        cell = ws2.cell(row=row, column=col,
                        value=f"=IF(Equity<>0,{cl}{ncf_row}/Equity,0)")
        cell.number_format = PERCENT_FMT
        cell.border = THIN_BORDER
        cell.fill = METRIC_FILL
    # Average CoC in total column
    ws2.cell(row=row, column=total_col,
             value=f"=IF(Equity<>0,{get_column_letter(total_col)}{ncf_row}/(Equity*{num_years}),0)").number_format = PERCENT_FMT
    ws2.cell(row=row, column=total_col).border = THIN_BORDER
    ws2.cell(row=row, column=total_col).fill = METRIC_FILL
    row += 1

    # DSCR row (formula = NOI / ABS(DS))
    dscr_cf_row = row
    ws2.cell(row=row, column=1, value="DSCR").font = METRIC_FONT
    ws2.cell(row=row, column=1).border = THIN_BORDER
    ws2.cell(row=row, column=1).fill = METRIC_FILL
    for i in range(num_years):
        col = i + 2
        cl = get_column_letter(col)
        cell = ws2.cell(row=row, column=col,
                        value=f"=IF(ABS({cl}{ds_row})<>0,{cl}{noi_row}/ABS({cl}{ds_row}),0)")
        cell.number_format = RATIO_FMT
        cell.border = THIN_BORDER
        cell.fill = METRIC_FILL
    ws2.cell(row=row, column=total_col).border = THIN_BORDER
    ws2.cell(row=row, column=total_col).fill = METRIC_FILL
    row += 1

    # Debt Yield row (formula = NOI / LoanAmount)
    dy_row = row
    ws2.cell(row=row, column=1, value="Debt Yield").font = METRIC_FONT
    ws2.cell(row=row, column=1).border = THIN_BORDER
    ws2.cell(row=row, column=1).fill = METRIC_FILL
    for i in range(num_years):
        col = i + 2
        cl = get_column_letter(col)
        cell = ws2.cell(row=row, column=col,
                        value=f"=IF(LoanAmount<>0,{cl}{noi_row}/LoanAmount,0)")
        cell.number_format = PERCENT_FMT
        cell.border = THIN_BORDER
        cell.fill = METRIC_FILL
    ws2.cell(row=row, column=total_col).border = THIN_BORDER
    ws2.cell(row=row, column=total_col).fill = METRIC_FILL
    row += 1

    # Loan Balance row
    lb_row = row
    ws2.cell(row=row, column=1, value="Loan Balance").font = METRIC_FONT
    ws2.cell(row=row, column=1).border = THIN_BORDER
    ws2.cell(row=row, column=1).fill = METRIC_FILL
    for i, y in enumerate(years):
        col = i + 2
        cell = ws2.cell(row=row, column=col, value=y.get("loanBal", 0) or 0)
        cell.number_format = CURRENCY_FMT
        cell.border = THIN_BORDER
        cell.fill = METRIC_FILL
    ws2.cell(row=row, column=total_col).border = THIN_BORDER
    ws2.cell(row=row, column=total_col).fill = METRIC_FILL
    row += 1

    # ══════════════════════════════════════════
    # Sheet 3: Returns (with formulas + Deal Summary)
    # ══════════════════════════════════════════
    ws3 = wb.create_sheet("Returns")
    ws3.column_dimensions["A"].width = 30
    ws3.column_dimensions["B"].width = 22

    ws3.merge_cells("A1:B1")
    ws3.cell(row=1, column=1, value="Investment Returns Summary").font = TITLE_FONT

    row = 3
    row = _section_header(ws3, row, "Core Metrics", 2)

    row = _add_kv_row(ws3, row, "Levered IRR", calc.get("levIRR", 0) / 100 if calc.get("levIRR") else 0, PERCENT_FMT)
    row = _add_kv_row(ws3, row, "Unlevered IRR", calc.get("unlevIRR", 0) / 100 if calc.get("unlevIRR") else 0, PERCENT_FMT)
    row = _add_kv_row(ws3, row, "Equity Multiple", calc.get("em", 0), RATIO_FMT)
    row = _add_kv_row(ws3, row, "Average CoC", calc.get("avgCoC", 0) / 100 if calc.get("avgCoC") else 0, PERCENT_FMT)
    row = _add_kv_row(ws3, row, "DSCR (Year 1)", calc.get("dscr", 0), RATIO_FMT)

    row += 1
    row = _section_header(ws3, row, "Capital Efficiency", 2)
    row = _add_kv_row(ws3, row, "Going-in Cap Rate", calc.get("goingCap", 0) / 100 if calc.get("goingCap") else 0, PERCENT_FMT)
    row = _add_kv_row(ws3, row, "Yield on Cost", calc.get("yoc", 0) / 100 if calc.get("yoc") else 0, PERCENT_FMT)

    row += 1
    row = _section_header(ws3, row, "Exit Analysis", 2)
    row = _add_kv_row(ws3, row, "Exit NOI", calc.get("exitNOI", 0), CURRENCY_FMT)
    row = _add_kv_row(ws3, row, "Exit Value", calc.get("exitVal", 0), CURRENCY_FMT)
    row = _add_kv_row(ws3, row, "Loan Payoff", calc.get("exitBal", 0), CURRENCY_FMT)
    row = _add_kv_row(ws3, row, "Net Sale Proceeds", calc.get("saleNet", 0), CURRENCY_FMT)

    row += 1
    row = _section_header(ws3, row, "Investment Decision", 2)
    go_cell = ws3.cell(row=row, column=1, value="GO / NO-GO")
    go_cell.font = Font(name="Calibri", bold=True, size=12)
    go_cell.border = THIN_BORDER
    verdict = "GO" if calc.get("goGreen") else "NO-GO"
    v_cell = ws3.cell(row=row, column=2, value=verdict)
    v_cell.font = Font(name="Calibri", bold=True, size=12, color="166534" if calc.get("goGreen") else "991B1B")
    v_cell.fill = GREEN_FILL if calc.get("goGreen") else RED_FILL
    v_cell.border = THIN_BORDER
    row += 2

    # ── Deal Summary Write-Up ──
    row = _section_header(ws3, row, "Executive Summary", 2)

    # Build summary paragraph
    sf_val = total_sf
    asset_type = p.get("assetType", "commercial")
    address = p.get("address", "N/A")
    irr_val = calc.get("levIRR", 0) or 0
    em_val = calc.get("em", 0) or 0
    y1_noi = calc.get("y1NOI", 0) or p.get("y1NOI", 0) or 0
    going_cap = calc.get("goingCap", 0) or 0
    price_psf = (purchase_price / sf_val) if sf_val > 0 else 0

    # Lease concentration note
    lease_note = ""
    if tenant_analytics and tenant_analytics.get("top_tenant_pct", 0) > 0:
        top_pct = tenant_analytics["top_tenant_pct"]
        if top_pct > 0.30:
            top_name = tenant_analytics["tenants_with_rev"][0].get("name", "top tenant") if tenant_analytics["tenants_with_rev"] else "top tenant"
            lease_note = f"concentration risk ({top_name} represents {top_pct*100:.0f}% of revenue)"
        else:
            lease_note = "diversified tenant base"
    else:
        lease_note = "tenant roll analysis pending"

    # Near-term expiry note
    expiry_note = ""
    if tenant_analytics and tenant_analytics.get("walt", 0) > 0:
        walt = tenant_analytics["walt"]
        if walt < 3:
            expiry_note = f"near-term lease rollover risk (WALT: {walt:.1f} years)"
        else:
            expiry_note = f"stable lease term (WALT: {walt:.1f} years)"
    else:
        expiry_note = "lease term analysis pending"

    summary_text = (
        f"The {prop_name} is a {sf_val:,.0f} SF {asset_type} property "
        f"located at {address}, offered at {_fmt_currency(purchase_price)} "
        f"({_fmt_currency(price_psf)}/SF). The going-in cap rate is {going_cap:.2f}% "
        f"with a projected levered IRR of {irr_val:.2f}% and {em_val:.2f}x equity multiple "
        f"over a {hold}-year hold period. The property generates Year 1 NOI of {_fmt_currency(y1_noi)}. "
        f"Key considerations include {lease_note} and {expiry_note}."
    )

    ws3.merge_cells(start_row=row, start_column=1, end_row=row + 3, end_column=2)
    summary_cell = ws3.cell(row=row, column=1, value=summary_text)
    summary_cell.font = Font(name="Calibri", size=10, italic=True)
    summary_cell.alignment = Alignment(wrap_text=True, vertical="top")
    summary_cell.border = THIN_BORDER

    # ══════════════════════════════════════════
    # Sheet 4: Waterfall (with tier labels)
    # ══════════════════════════════════════════
    ws4 = wb.create_sheet("Waterfall")
    ws4.column_dimensions["A"].width = 28
    ws4.column_dimensions["B"].width = 20
    ws4.column_dimensions["C"].width = 20

    ws4.merge_cells("A1:C1")
    ws4.cell(row=1, column=1, value="LP / GP Waterfall Distribution").font = TITLE_FONT

    row = 3
    row = _section_header(ws4, row, "Waterfall Configuration", 3)
    row = _add_kv_row(ws4, row, "LP Equity %", (wf.get("lpPercent", 90) or 90) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "GP Equity %", (wf.get("gpPercent", 10) or 10) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "Preferred Return", (wf.get("prefReturn", 8) or 8) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "Catch-Up", "Yes" if wf.get("catchUp") else "No")
    row = _add_kv_row(ws4, row, "Tier 1 IRR Threshold", (wf.get("tier1Thresh", 15) or 15) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "Tier 1 LP Split", (wf.get("tier1Split", 80) or 80) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "Tier 2 IRR Threshold", (wf.get("tier2Thresh", 20) or 20) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "Tier 2 LP Split", (wf.get("tier2Split", 70) or 70) / 100, PERCENT_FMT)
    row = _add_kv_row(ws4, row, "Tier 3 LP Split", (wf.get("tier3Split", 60) or 60) / 100, PERCENT_FMT)

    row += 1
    row = _section_header(ws4, row, "Distribution Summary", 3)
    # Header row
    for ci, h in enumerate(["Metric", "LP", "GP"], 1):
        cell = ws4.cell(row=row, column=ci, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER
    row += 1

    dist_rows = [
        ("Equity Contribution", calc.get("lpEq", 0), calc.get("gpEq", 0)),
        ("Total Distributions", calc.get("lpOut", 0), calc.get("gpOut", 0)),
        ("Net Profit", (calc.get("lpOut", 0) or 0) - (calc.get("lpEq", 0) or 0),
         (calc.get("gpOut", 0) or 0) - (calc.get("gpEq", 0) or 0)),
    ]
    for label, lp_val, gp_val in dist_rows:
        ws4.cell(row=row, column=1, value=label).font = LABEL_FONT
        ws4.cell(row=row, column=1).border = THIN_BORDER
        ws4.cell(row=row, column=2, value=lp_val).number_format = CURRENCY_FMT
        ws4.cell(row=row, column=2).border = THIN_BORDER
        ws4.cell(row=row, column=3, value=gp_val).number_format = CURRENCY_FMT
        ws4.cell(row=row, column=3).border = THIN_BORDER
        row += 1

    # IRR and EM
    lp_irr = calc.get("lpIRR")
    gp_irr = calc.get("gpIRR")
    ws4.cell(row=row, column=1, value="IRR").font = LABEL_FONT
    ws4.cell(row=row, column=1).border = THIN_BORDER
    ws4.cell(row=row, column=2, value=lp_irr / 100 if lp_irr else 0).number_format = PERCENT_FMT
    ws4.cell(row=row, column=2).border = THIN_BORDER
    ws4.cell(row=row, column=3, value=gp_irr / 100 if gp_irr else 0).number_format = PERCENT_FMT
    ws4.cell(row=row, column=3).border = THIN_BORDER
    row += 1

    ws4.cell(row=row, column=1, value="Equity Multiple").font = LABEL_FONT
    ws4.cell(row=row, column=1).border = THIN_BORDER
    ws4.cell(row=row, column=2, value=calc.get("lpEM", 0)).number_format = RATIO_FMT
    ws4.cell(row=row, column=2).border = THIN_BORDER
    gp_eq = calc.get("gpEq", 0) or 1
    ws4.cell(row=row, column=3, value=(calc.get("gpOut", 0) or 0) / gp_eq).number_format = RATIO_FMT
    ws4.cell(row=row, column=3).border = THIN_BORDER
    row += 1

    ws4.cell(row=row, column=1, value="GP Promote").font = Font(name="Calibri", bold=True, size=11, color="166534")
    ws4.cell(row=row, column=1).border = THIN_BORDER
    ws4.cell(row=row, column=2).border = THIN_BORDER
    promote_cell = ws4.cell(row=row, column=3, value=calc.get("gpPromote", 0))
    promote_cell.number_format = CURRENCY_FMT
    promote_cell.font = Font(name="Calibri", bold=True, size=11, color="166534")
    promote_cell.fill = GREEN_FILL
    promote_cell.border = THIN_BORDER

    # ══════════════════════════════════════════
    # Sheet 5: Value-Add & CapEx (keep existing)
    # ══════════════════════════════════════════
    if events or cap_items:
        ws5 = wb.create_sheet("Value-Add & CapEx")
        ws5.column_dimensions["A"].width = 28
        ws5.column_dimensions["B"].width = 12
        ws5.column_dimensions["C"].width = 16
        ws5.column_dimensions["D"].width = 14

        ws5.merge_cells("A1:D1")
        ws5.cell(row=1, column=1, value="Value-Add Events & Capital Expenditures").font = TITLE_FONT

        row = 3
        if events:
            row = _section_header(ws5, row, "Value-Add Events", 4)
            for ci, h in enumerate(["Label", "Year", "Amount", "Type"], 1):
                cell = ws5.cell(row=row, column=ci, value=h)
                cell.font = HEADER_FONT
                cell.fill = HEADER_FILL
                cell.alignment = Alignment(horizontal="center")
                cell.border = THIN_BORDER
            row += 1
            for ev in events:
                ws5.cell(row=row, column=1, value=ev.get("label", "")).border = THIN_BORDER
                ws5.cell(row=row, column=2, value=ev.get("year", 0)).border = THIN_BORDER
                ws5.cell(row=row, column=3, value=ev.get("amount", 0)).number_format = CURRENCY_FMT
                ws5.cell(row=row, column=3).border = THIN_BORDER
                etype = ev.get("type", "cost")
                t_cell = ws5.cell(row=row, column=4, value=etype.title())
                t_cell.border = THIN_BORDER
                t_cell.font = GREEN_FONT if etype == "income" else RED_FONT
                row += 1
            row += 1

        if cap_items:
            row = _section_header(ws5, row, "CapEx Items", 3)
            for ci, h in enumerate(["Label", "Year", "Amount"], 1):
                cell = ws5.cell(row=row, column=ci, value=h)
                cell.font = HEADER_FONT
                cell.fill = HEADER_FILL
                cell.alignment = Alignment(horizontal="center")
                cell.border = THIN_BORDER
            row += 1
            total_capex = 0
            for ci_item in cap_items:
                ws5.cell(row=row, column=1, value=ci_item.get("label", "")).border = THIN_BORDER
                ws5.cell(row=row, column=2, value=ci_item.get("year", 0)).border = THIN_BORDER
                amt = ci_item.get("amount", 0)
                ws5.cell(row=row, column=3, value=amt).number_format = CURRENCY_FMT
                ws5.cell(row=row, column=3).border = THIN_BORDER
                total_capex += amt
                row += 1
            ws5.cell(row=row, column=1, value="TOTAL").font = Font(name="Calibri", bold=True, size=10)
            ws5.cell(row=row, column=1).border = THIN_BORDER
            ws5.cell(row=row, column=3, value=total_capex).number_format = CURRENCY_FMT
            ws5.cell(row=row, column=3).font = Font(name="Calibri", bold=True)
            ws5.cell(row=row, column=3).border = THIN_BORDER

    # ══════════════════════════════════════════
    # Sheet 6: Rent Roll
    # ══════════════════════════════════════════
    if tenants:
        ws6 = wb.create_sheet("Rent Roll")
        rr_h = ["Tenant", "Suite", "SF", "Type", "Rent/SF", "CAM/SF", "Escal %", "Start", "End", "TI/SF", "LC %"]
        rr_w = [24, 12, 12, 10, 12, 12, 10, 14, 14, 10, 10]
        for i, w in enumerate(rr_w):
            ws6.column_dimensions[chr(65 + i)].width = w

        for ci, h in enumerate(rr_h, 1):
            cell = ws6.cell(row=1, column=ci, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER

        for ri, t in enumerate(tenants, 2):
            ws6.cell(row=ri, column=1, value=t.get("name", "")).border = THIN_BORDER
            ws6.cell(row=ri, column=2, value=t.get("suite", "")).border = THIN_BORDER
            ws6.cell(row=ri, column=3, value=t.get("sf", 0)).number_format = NUMBER_FMT
            ws6.cell(row=ri, column=3).border = THIN_BORDER
            ws6.cell(row=ri, column=4, value=t.get("type", "NNN")).border = THIN_BORDER
            ws6.cell(row=ri, column=5, value=t.get("rentPSF", 0)).number_format = CURRENCY_CENTS_FMT
            ws6.cell(row=ri, column=5).border = THIN_BORDER
            ws6.cell(row=ri, column=6, value=t.get("camPSF", 0)).number_format = CURRENCY_CENTS_FMT
            ws6.cell(row=ri, column=6).border = THIN_BORDER
            ws6.cell(row=ri, column=7, value=(t.get("escalPct", 0) or 0) / 100).number_format = PERCENT_FMT
            ws6.cell(row=ri, column=7).border = THIN_BORDER
            ws6.cell(row=ri, column=8, value=t.get("start", "")).border = THIN_BORDER
            ws6.cell(row=ri, column=9, value=t.get("end", "")).border = THIN_BORDER
            ws6.cell(row=ri, column=10, value=t.get("tiPSF", 0)).number_format = CURRENCY_CENTS_FMT
            ws6.cell(row=ri, column=10).border = THIN_BORDER
            ws6.cell(row=ri, column=11, value=(t.get("lcPct", 0) or 0) / 100).number_format = PERCENT_FMT
            ws6.cell(row=ri, column=11).border = THIN_BORDER

    # ══════════════════════════════════════════
    # Sheet 7: Key Insights
    # ══════════════════════════════════════════
    ws_insights = wb.create_sheet("Key Insights")
    ws_insights.column_dimensions["A"].width = 28
    ws_insights.column_dimensions["B"].width = 18
    ws_insights.column_dimensions["C"].width = 18
    ws_insights.column_dimensions["D"].width = 18
    ws_insights.column_dimensions["E"].width = 18

    ws_insights.merge_cells("A1:E1")
    ws_insights.cell(row=1, column=1, value="Key Insights & Tenant Analytics").font = TITLE_FONT

    row = 3

    if tenants and tenant_analytics:
        # ── Lease Expiration Profile ──
        row = _section_header(ws_insights, row, "Lease Expiration Profile", 5)
        for ci, h in enumerate(["Year", "# Tenants Expiring", "SF Expiring", "% of Total SF", "Revenue At Risk"], 1):
            cell = ws_insights.cell(row=row, column=ci, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER
        row += 1

        expiry_buckets = tenant_analytics.get("expiry_buckets", {})
        for yr_key in sorted(expiry_buckets.keys()):
            bucket_tenants = expiry_buckets[yr_key]
            n_tenants = len(bucket_tenants)
            sf_expiring = sum(t.get("sf", 0) or 0 for t in bucket_tenants)
            pct_sf = sf_expiring / total_sf if total_sf > 0 else 0
            rev_at_risk = sum((t.get("sf", 0) or 0) * (t.get("rentPSF", 0) or 0) for t in bucket_tenants)

            ws_insights.cell(row=row, column=1, value=yr_key).border = THIN_BORDER
            ws_insights.cell(row=row, column=2, value=n_tenants).border = THIN_BORDER
            ws_insights.cell(row=row, column=2).alignment = Alignment(horizontal="center")
            ws_insights.cell(row=row, column=3, value=sf_expiring).number_format = NUMBER_FMT
            ws_insights.cell(row=row, column=3).border = THIN_BORDER
            pct_cell = ws_insights.cell(row=row, column=4, value=pct_sf)
            pct_cell.number_format = PERCENT_FMT
            pct_cell.border = THIN_BORDER
            if pct_sf > 0.25:
                pct_cell.fill = RED_FILL
                pct_cell.font = RED_FONT
            elif pct_sf > 0.15:
                pct_cell.fill = AMBER_FILL
            ws_insights.cell(row=row, column=5, value=rev_at_risk).number_format = CURRENCY_FMT
            ws_insights.cell(row=row, column=5).border = THIN_BORDER
            row += 1

        row += 1

        # ── Top 5 Tenants by Revenue ──
        row = _section_header(ws_insights, row, "Top 5 Tenants by Revenue", 5)
        for ci, h in enumerate(["Tenant", "SF", "Rent/SF", "Annual Revenue", "% of Total"], 1):
            cell = ws_insights.cell(row=row, column=ci, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER
        row += 1

        top5 = tenant_analytics.get("tenants_with_rev", [])[:5]
        total_rev = tenant_analytics.get("total_revenue", 0)
        for t in top5:
            ws_insights.cell(row=row, column=1, value=t.get("name", "")).border = THIN_BORDER
            ws_insights.cell(row=row, column=2, value=t.get("_sf", 0)).number_format = NUMBER_FMT
            ws_insights.cell(row=row, column=2).border = THIN_BORDER
            ws_insights.cell(row=row, column=3, value=t.get("rentPSF", 0)).number_format = CURRENCY_CENTS_FMT
            ws_insights.cell(row=row, column=3).border = THIN_BORDER
            ws_insights.cell(row=row, column=4, value=t.get("_annual_rev", 0)).number_format = CURRENCY_FMT
            ws_insights.cell(row=row, column=4).border = THIN_BORDER
            pct = t.get("_annual_rev", 0) / total_rev if total_rev > 0 else 0
            ws_insights.cell(row=row, column=5, value=pct).number_format = PERCENT_FMT
            ws_insights.cell(row=row, column=5).border = THIN_BORDER
            row += 1

        row += 1

        # ── Concentration & WALT ──
        row = _section_header(ws_insights, row, "Risk Metrics", 2)
        top_pct = tenant_analytics.get("top_tenant_pct", 0)
        top_name = top5[0].get("name", "N/A") if top5 else "N/A"
        row = _add_kv_row(ws_insights, row, f"Top Tenant ({top_name})", top_pct, PERCENT_FMT)
        walt = tenant_analytics.get("walt", 0)
        row = _add_kv_row(ws_insights, row, "Weighted Avg Lease Term (WALT)", f"{walt:.1f} years")

        # Color-code concentration
        conc_cell = ws_insights.cell(row=row, column=1, value="Concentration Risk")
        conc_cell.font = LABEL_FONT
        conc_cell.border = THIN_BORDER
        if top_pct > 0.30:
            risk_cell = ws_insights.cell(row=row, column=2, value="HIGH")
            risk_cell.fill = RED_FILL
            risk_cell.font = RED_FONT
        elif top_pct > 0.15:
            risk_cell = ws_insights.cell(row=row, column=2, value="MODERATE")
            risk_cell.fill = AMBER_FILL
        else:
            risk_cell = ws_insights.cell(row=row, column=2, value="LOW")
            risk_cell.fill = GREEN_FILL
            risk_cell.font = GREEN_FONT
        risk_cell.border = THIN_BORDER
    else:
        ws_insights.cell(row=row, column=1, value="No tenant data available. Add tenants in the Rent Roll tab to generate analytics.").font = Font(name="Calibri", italic=True, size=10)

    # ══════════════════════════════════════════
    # Sheet 8: Charts
    # ══════════════════════════════════════════
    if num_years > 0:
        ws_charts = wb.create_sheet("Charts")
        ws_charts.column_dimensions["A"].width = 10

        # ── Data table for charts (hidden-ish at top of Charts sheet) ──
        chart_data_start = 1
        ws_charts.cell(row=1, column=1, value="Year").font = HEADER_FONT
        ws_charts.cell(row=1, column=1).fill = HEADER_FILL
        ws_charts.cell(row=1, column=2, value="NOI").font = HEADER_FONT
        ws_charts.cell(row=1, column=2).fill = HEADER_FILL
        ws_charts.cell(row=1, column=3, value="Debt Service").font = HEADER_FONT
        ws_charts.cell(row=1, column=3).fill = HEADER_FILL
        ws_charts.cell(row=1, column=4, value="Net Cash Flow").font = HEADER_FONT
        ws_charts.cell(row=1, column=4).fill = HEADER_FILL
        ws_charts.cell(row=1, column=5, value="Loan Balance").font = HEADER_FONT
        ws_charts.cell(row=1, column=5).fill = HEADER_FILL

        for i, y in enumerate(years):
            r = i + 2
            ws_charts.cell(row=r, column=1, value=f"Year {y.get('yr', i+1)}")
            ws_charts.cell(row=r, column=2, value=y.get("noi", 0))
            ws_charts.cell(row=r, column=3, value=abs(y.get("annDS", 0) or 0))
            ncf_val = (y.get("noi", 0) or 0) - abs(y.get("annDS", 0) or 0) - (y.get("capexRes", 0) or 0) - (y.get("specCapex", 0) or 0) - (y.get("tiLC", 0) or 0) + (y.get("vaInc", 0) or 0) - (y.get("vaCost", 0) or 0)
            ws_charts.cell(row=r, column=4, value=ncf_val)
            ws_charts.cell(row=r, column=5, value=y.get("loanBal", 0) or 0)

        data_end = num_years + 1

        # ── Chart A: NOI vs DS vs NCF (Bar) ──
        chart1 = BarChart()
        chart1.type = "col"
        chart1.style = 10
        chart1.title = "NOI vs Debt Service vs Net Cash Flow"
        chart1.y_axis.title = "Amount ($)"
        chart1.x_axis.title = "Year"
        chart1.width = 20
        chart1.height = 12

        cats = Reference(ws_charts, min_col=1, min_row=2, max_row=data_end)
        noi_data = Reference(ws_charts, min_col=2, min_row=1, max_row=data_end)
        ds_data = Reference(ws_charts, min_col=3, min_row=1, max_row=data_end)
        ncf_data = Reference(ws_charts, min_col=4, min_row=1, max_row=data_end)

        chart1.add_data(noi_data, titles_from_data=True)
        chart1.add_data(ds_data, titles_from_data=True)
        chart1.add_data(ncf_data, titles_from_data=True)
        chart1.set_categories(cats)

        # Color the series
        chart1.series[0].graphicalProperties.solidFill = "22C55E"  # green for NOI
        chart1.series[1].graphicalProperties.solidFill = "EF4444"  # red for DS
        chart1.series[2].graphicalProperties.solidFill = "3B82F6"  # blue for NCF

        ws_charts.add_chart(chart1, f"A{data_end + 3}")

        # ── Chart B: Loan Balance (Line) ──
        chart2 = LineChart()
        chart2.style = 10
        chart2.title = "Loan Balance Over Hold Period"
        chart2.y_axis.title = "Loan Balance ($)"
        chart2.x_axis.title = "Year"
        chart2.width = 20
        chart2.height = 12

        lb_data = Reference(ws_charts, min_col=5, min_row=1, max_row=data_end)
        chart2.add_data(lb_data, titles_from_data=True)
        chart2.set_categories(cats)
        chart2.series[0].graphicalProperties.line.solidFill = "F59E0B"  # amber

        ws_charts.add_chart(chart2, f"A{data_end + 20}")

        # ── Chart C: Lease Expiration Profile (Bar) — if tenants exist ──
        if tenants and tenant_analytics and tenant_analytics.get("expiry_buckets"):
            expiry_buckets = tenant_analytics["expiry_buckets"]
            sorted_years = sorted(expiry_buckets.keys())

            # Write expiry data for chart
            exp_start_row = data_end + 38
            ws_charts.cell(row=exp_start_row, column=1, value="Expiry Year").font = HEADER_FONT
            ws_charts.cell(row=exp_start_row, column=1).fill = HEADER_FILL
            ws_charts.cell(row=exp_start_row, column=2, value="# Tenants").font = HEADER_FONT
            ws_charts.cell(row=exp_start_row, column=2).fill = HEADER_FILL
            ws_charts.cell(row=exp_start_row, column=3, value="SF Expiring").font = HEADER_FONT
            ws_charts.cell(row=exp_start_row, column=3).fill = HEADER_FILL

            for idx, yr_key in enumerate(sorted_years):
                r = exp_start_row + 1 + idx
                ws_charts.cell(row=r, column=1, value=yr_key)
                ws_charts.cell(row=r, column=2, value=len(expiry_buckets[yr_key]))
                sf_exp = sum(t.get("sf", 0) or 0 for t in expiry_buckets[yr_key])
                ws_charts.cell(row=r, column=3, value=sf_exp)

            exp_end = exp_start_row + len(sorted_years)

            chart3 = BarChart()
            chart3.type = "col"
            chart3.style = 10
            chart3.title = "Lease Expiration Profile"
            chart3.y_axis.title = "Count / SF"
            chart3.x_axis.title = "Year"
            chart3.width = 20
            chart3.height = 12

            exp_cats = Reference(ws_charts, min_col=1, min_row=exp_start_row + 1, max_row=exp_end)
            exp_tenant_data = Reference(ws_charts, min_col=2, min_row=exp_start_row, max_row=exp_end)
            exp_sf_data = Reference(ws_charts, min_col=3, min_row=exp_start_row, max_row=exp_end)

            chart3.add_data(exp_tenant_data, titles_from_data=True)
            chart3.add_data(exp_sf_data, titles_from_data=True)
            chart3.set_categories(exp_cats)

            chart3.series[0].graphicalProperties.solidFill = "8B5CF6"  # purple
            chart3.series[1].graphicalProperties.solidFill = "F97316"  # orange

            ws_charts.add_chart(chart3, f"A{exp_end + 3}")

    # ── Write to buffer ──
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    safe_name = prop_name.replace(" ", "_").replace("/", "-")
    safe_name = safe_name.encode("ascii", "ignore").decode("ascii")[:40]
    filename = f"{safe_name}_V2_Analysis.xlsx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════════════════════
# V2 MEMO HTML EXPORT
# ═══════════════════════════════════════════════════════════

@router.post("/v2/memo/html")
async def export_v2_memo_html(data: V2ExportRequest):
    """Generate a styled HTML memo for the deal — can be printed to PDF via browser."""
    state = data.v2_state or {}
    calc = data.calc or {}
    p = state.get("assumptions", {})
    wf = state.get("waterfall", {})
    tenants = state.get("tenants", [])
    years = calc.get("years", [])

    prop_name = p.get("name", "Untitled Deal")
    hold = p.get("holdPeriod", 5)
    total_sf = p.get("sf", 0) or 0
    purchase_price = p.get("purchasePrice", calc.get("pp", 0)) or 0
    ltv_pct = (p.get("ltv", 65) or 65)
    loan_amount = purchase_price * ltv_pct / 100
    equity_val = calc.get("totalEq", 0) or calc.get("equity", 0) or (purchase_price - loan_amount)
    irr_val = calc.get("levIRR", 0) or 0
    em_val = calc.get("em", 0) or 0
    going_cap = calc.get("goingCap", 0) or 0
    y1_noi = calc.get("y1NOI", 0) or p.get("y1NOI", 0) or 0
    price_psf = purchase_price / total_sf if total_sf > 0 else 0

    tenant_analytics = _v2_build_tenant_analytics(tenants, total_sf)

    # Build HTML
    html_parts = []
    html_parts.append(f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{prop_name} — Investment Memo</title>
<style>
  @page {{ size: letter; margin: 1in; }}
  body {{ font-family: 'Segoe UI', Calibri, Arial, sans-serif; color: #1a1a2e; line-height: 1.6; max-width: 900px; margin: 0 auto; padding: 40px; }}
  h1 {{ color: #0A1628; border-bottom: 3px solid #0A1628; padding-bottom: 10px; }}
  h2 {{ color: #132A42; border-bottom: 1px solid #ccc; padding-bottom: 5px; margin-top: 30px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
  th {{ background: #0A1628; color: white; padding: 8px 12px; text-align: left; }}
  td {{ padding: 6px 12px; border-bottom: 1px solid #e0e0e0; }}
  tr:nth-child(even) {{ background: #f8f9fa; }}
  .metric {{ display: inline-block; background: #EFF6FF; border: 1px solid #DBEAFE; border-radius: 8px; padding: 12px 20px; margin: 5px; text-align: center; min-width: 140px; }}
  .metric .label {{ font-size: 11px; color: #6b7280; text-transform: uppercase; }}
  .metric .value {{ font-size: 20px; font-weight: bold; color: #0A1628; }}
  .go {{ color: #166534; background: #DCFCE7; padding: 4px 12px; border-radius: 4px; font-weight: bold; }}
  .nogo {{ color: #991B1B; background: #FEE2E2; padding: 4px 12px; border-radius: 4px; font-weight: bold; }}
  .summary {{ background: #f0f4f8; border-left: 4px solid #0A1628; padding: 15px 20px; margin: 20px 0; font-style: italic; }}
  .footer {{ margin-top: 40px; padding-top: 15px; border-top: 1px solid #ccc; font-size: 11px; color: #999; }}
</style>
</head>
<body>
<h1>{prop_name}</h1>
<p style="color: #6b7280; margin-top: -10px;">Investment Memo &mdash; Generated {datetime.now().strftime('%B %d, %Y')}</p>
""")

    # Executive Summary
    lease_note = "diversified tenant base"
    expiry_note = "lease term analysis pending"
    if tenant_analytics:
        top_pct = tenant_analytics.get("top_tenant_pct", 0)
        walt = tenant_analytics.get("walt", 0)
        if top_pct > 0.30:
            top_name = tenant_analytics["tenants_with_rev"][0].get("name", "top tenant") if tenant_analytics.get("tenants_with_rev") else "top tenant"
            lease_note = f"concentration risk ({top_name} at {top_pct*100:.0f}% of revenue)"
        if walt > 0:
            expiry_note = f"WALT of {walt:.1f} years"

    html_parts.append(f"""
<h2>Executive Summary</h2>
<div class="summary">
The {prop_name} is a {total_sf:,.0f} SF {p.get('assetType', 'commercial')} property
located at {p.get('address', 'N/A')}, offered at {_fmt_currency(purchase_price)}
({_fmt_currency(price_psf)}/SF). The going-in cap rate is {going_cap:.2f}%
with a projected levered IRR of {irr_val:.2f}% and {em_val:.2f}x equity multiple
over a {hold}-year hold period. Year 1 NOI: {_fmt_currency(y1_noi)}.
Key considerations include {lease_note} and {expiry_note}.
</div>
""")

    # Key Metrics
    verdict_class = "go" if calc.get("goGreen") else "nogo"
    verdict_text = "GO" if calc.get("goGreen") else "NO-GO"
    html_parts.append(f"""
<h2>Key Return Metrics</h2>
<div style="display: flex; flex-wrap: wrap; gap: 10px; margin: 15px 0;">
  <div class="metric"><div class="label">Levered IRR</div><div class="value">{irr_val:.2f}%</div></div>
  <div class="metric"><div class="label">Equity Multiple</div><div class="value">{em_val:.2f}x</div></div>
  <div class="metric"><div class="label">Avg CoC</div><div class="value">{(calc.get('avgCoC', 0) or 0):.2f}%</div></div>
  <div class="metric"><div class="label">DSCR</div><div class="value">{(calc.get('dscr', 0) or 0):.2f}x</div></div>
  <div class="metric"><div class="label">Going-In Cap</div><div class="value">{going_cap:.2f}%</div></div>
  <div class="metric"><div class="label">Verdict</div><div class="value"><span class="{verdict_class}">{verdict_text}</span></div></div>
</div>
""")

    # Capital Structure
    html_parts.append(f"""
<h2>Capital Structure</h2>
<table>
  <tr><th>Item</th><th>Amount</th><th>Per SF</th></tr>
  <tr><td>Purchase Price</td><td>{_fmt_currency(purchase_price)}</td><td>{_fmt_currency(price_psf)}</td></tr>
  <tr><td>Loan Amount ({ltv_pct:.0f}% LTV)</td><td>{_fmt_currency(loan_amount)}</td><td>{_fmt_currency(loan_amount/total_sf if total_sf > 0 else 0)}</td></tr>
  <tr><td>Equity Required</td><td>{_fmt_currency(equity_val)}</td><td>{_fmt_currency(equity_val/total_sf if total_sf > 0 else 0)}</td></tr>
</table>
""")

    # Cash Flow table
    if years:
        html_parts.append("<h2>Cash Flow Projection</h2><table><tr><th>Year</th><th>NOI</th><th>Debt Service</th><th>Net CF</th></tr>")
        for y in years:
            noi_v = y.get("noi", 0) or 0
            ds_v = y.get("annDS", 0) or 0
            ncf_v = noi_v - abs(ds_v) - (y.get("capexRes", 0) or 0) - (y.get("specCapex", 0) or 0) - (y.get("tiLC", 0) or 0)
            html_parts.append(f"<tr><td>Year {y.get('yr', '')}</td><td>{_fmt_currency(noi_v)}</td><td>{_fmt_currency(abs(ds_v))}</td><td>{_fmt_currency(ncf_v)}</td></tr>")
        html_parts.append("</table>")

    # Waterfall
    if wf:
        lp_out = calc.get("lpOut", 0) or 0
        gp_out = calc.get("gpOut", 0) or 0
        html_parts.append(f"""
<h2>Waterfall Distribution</h2>
<table>
  <tr><th>Metric</th><th>LP</th><th>GP</th></tr>
  <tr><td>Equity</td><td>{_fmt_currency(calc.get('lpEq', 0))}</td><td>{_fmt_currency(calc.get('gpEq', 0))}</td></tr>
  <tr><td>Total Distributions</td><td>{_fmt_currency(lp_out)}</td><td>{_fmt_currency(gp_out)}</td></tr>
  <tr><td>IRR</td><td>{(calc.get('lpIRR', 0) or 0):.2f}%</td><td>{(calc.get('gpIRR', 0) or 0):.2f}%</td></tr>
  <tr><td>GP Promote</td><td></td><td>{_fmt_currency(calc.get('gpPromote', 0))}</td></tr>
</table>
""")

    html_parts.append(f"""
<div class="footer">
  <p>Generated by CRE Lytic &mdash; {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
  <p>This memo is for informational purposes only and does not constitute investment advice.</p>
</div>
</body></html>""")

    html_content = "".join(html_parts)

    safe_name = prop_name.replace(" ", "_").replace("/", "-")
    safe_name = safe_name.encode("ascii", "ignore").decode("ascii")[:40]
    filename = f"{safe_name}_Memo.html"

    return StreamingResponse(
        io.BytesIO(html_content.encode("utf-8")),
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════════════════════
# V2 MEMO DOCX EXPORT
# ═══════════════════════════════════════════════════════════

@router.post("/v2/memo/docx")
async def export_v2_memo_docx(data: V2ExportRequest):
    """Generate a professional Word memo from V2 deal data."""
    if not HAS_DOCX:
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=501,
            content={"detail": "python-docx is not installed. Run: pip install python-docx"},
        )

    state = data.v2_state or {}
    calc = data.calc or {}
    p = state.get("assumptions", {})
    wf = state.get("waterfall", {})
    tenants = state.get("tenants", [])
    years = calc.get("years", [])

    prop_name = p.get("name", "Untitled Deal")
    hold = p.get("holdPeriod", 5)
    total_sf = p.get("sf", 0) or 0
    purchase_price = p.get("purchasePrice", calc.get("pp", 0)) or 0
    ltv_pct = (p.get("ltv", 65) or 65)
    loan_amount = purchase_price * ltv_pct / 100
    equity_val = calc.get("totalEq", 0) or calc.get("equity", 0) or (purchase_price - loan_amount)
    irr_val = calc.get("levIRR", 0) or 0
    em_val = calc.get("em", 0) or 0
    going_cap = calc.get("goingCap", 0) or 0
    y1_noi = calc.get("y1NOI", 0) or p.get("y1NOI", 0) or 0
    price_psf = purchase_price / total_sf if total_sf > 0 else 0

    tenant_analytics = _v2_build_tenant_analytics(tenants, total_sf)

    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Title
    title_para = doc.add_heading(prop_name, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Investment Memo — {datetime.now().strftime('%B %d, %Y')}")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = None  # default
    subtitle_run.italic = True

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)

    lease_note = "diversified tenant base"
    expiry_note = "lease term analysis pending"
    if tenant_analytics:
        top_pct = tenant_analytics.get("top_tenant_pct", 0)
        walt = tenant_analytics.get("walt", 0)
        if top_pct > 0.30:
            top_name = tenant_analytics["tenants_with_rev"][0].get("name", "top tenant") if tenant_analytics.get("tenants_with_rev") else "top tenant"
            lease_note = f"concentration risk ({top_name} at {top_pct*100:.0f}% of revenue)"
        if walt > 0:
            expiry_note = f"WALT of {walt:.1f} years"

    summary_text = (
        f"The {prop_name} is a {total_sf:,.0f} SF {p.get('assetType', 'commercial')} property "
        f"located at {p.get('address', 'N/A')}, offered at {_fmt_currency(purchase_price)} "
        f"({_fmt_currency(price_psf)}/SF). The going-in cap rate is {going_cap:.2f}% "
        f"with a projected levered IRR of {irr_val:.2f}% and {em_val:.2f}x equity multiple "
        f"over a {hold}-year hold period. The property generates Year 1 NOI of {_fmt_currency(y1_noi)}. "
        f"Key considerations include {lease_note} and {expiry_note}."
    )
    summary_para = doc.add_paragraph(summary_text)
    summary_para.style.font.size = Pt(10)

    # ── Capital Structure ──
    doc.add_heading("Capital Structure", level=1)
    cap_table = doc.add_table(rows=4, cols=3)
    cap_table.style = 'Light Grid Accent 1'
    cap_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cap_headers = ["Item", "Amount", "Per SF"]
    for i, h in enumerate(cap_headers):
        cap_table.rows[0].cells[i].text = h
        for paragraph in cap_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    cap_data = [
        ("Purchase Price", _fmt_currency(purchase_price), _fmt_currency(price_psf)),
        (f"Loan ({ltv_pct:.0f}% LTV)", _fmt_currency(loan_amount), _fmt_currency(loan_amount/total_sf if total_sf > 0 else 0)),
        ("Equity", _fmt_currency(equity_val), _fmt_currency(equity_val/total_sf if total_sf > 0 else 0)),
    ]
    for ri, (label, amt, psf) in enumerate(cap_data, 1):
        cap_table.rows[ri].cells[0].text = label
        cap_table.rows[ri].cells[1].text = str(amt)
        cap_table.rows[ri].cells[2].text = str(psf)

    # ── Returns Summary ──
    doc.add_heading("Returns Summary", level=1)
    ret_table = doc.add_table(rows=6, cols=2)
    ret_table.style = 'Light Grid Accent 1'
    ret_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    ret_data = [
        ("Metric", "Value"),
        ("Levered IRR", f"{irr_val:.2f}%"),
        ("Equity Multiple", f"{em_val:.2f}x"),
        ("Avg Cash-on-Cash", f"{(calc.get('avgCoC', 0) or 0):.2f}%"),
        ("DSCR (Year 1)", f"{(calc.get('dscr', 0) or 0):.2f}x"),
        ("Verdict", "GO" if calc.get("goGreen") else "NO-GO"),
    ]
    for ri, (label, val) in enumerate(ret_data):
        ret_table.rows[ri].cells[0].text = label
        ret_table.rows[ri].cells[1].text = str(val)
        if ri == 0:
            for cell in ret_table.rows[ri].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # ── Cash Flow Table ──
    if years:
        doc.add_heading("Cash Flow Projection", level=1)
        cf_table = doc.add_table(rows=len(years) + 1, cols=4)
        cf_table.style = 'Light Grid Accent 1'
        cf_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        cf_headers = ["Year", "NOI", "Debt Service", "Net CF"]
        for i, h in enumerate(cf_headers):
            cf_table.rows[0].cells[i].text = h
            for paragraph in cf_table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for ri, y in enumerate(years, 1):
            noi_v = y.get("noi", 0) or 0
            ds_v = abs(y.get("annDS", 0) or 0)
            ncf_v = noi_v - ds_v - (y.get("capexRes", 0) or 0) - (y.get("specCapex", 0) or 0) - (y.get("tiLC", 0) or 0)
            cf_table.rows[ri].cells[0].text = f"Year {y.get('yr', ri)}"
            cf_table.rows[ri].cells[1].text = _fmt_currency(noi_v)
            cf_table.rows[ri].cells[2].text = _fmt_currency(ds_v)
            cf_table.rows[ri].cells[3].text = _fmt_currency(ncf_v)

    # ── Waterfall Distribution ──
    if wf:
        doc.add_heading("Waterfall Distribution", level=1)
        wf_table = doc.add_table(rows=5, cols=3)
        wf_table.style = 'Light Grid Accent 1'
        wf_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        wf_data = [
            ("Metric", "LP", "GP"),
            ("Equity", _fmt_currency(calc.get("lpEq", 0)), _fmt_currency(calc.get("gpEq", 0))),
            ("Total Distributions", _fmt_currency(calc.get("lpOut", 0)), _fmt_currency(calc.get("gpOut", 0))),
            ("IRR", f"{(calc.get('lpIRR', 0) or 0):.2f}%", f"{(calc.get('gpIRR', 0) or 0):.2f}%"),
            ("GP Promote", "", _fmt_currency(calc.get("gpPromote", 0))),
        ]
        for ri, row_data in enumerate(wf_data):
            for ci, val in enumerate(row_data):
                wf_table.rows[ri].cells[ci].text = str(val)
            if ri == 0:
                for cell in wf_table.rows[ri].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # ── Key Insights ──
    if tenants and tenant_analytics:
        doc.add_heading("Key Insights", level=1)

        walt = tenant_analytics.get("walt", 0)
        top_pct = tenant_analytics.get("top_tenant_pct", 0)
        top5 = tenant_analytics.get("tenants_with_rev", [])[:5]

        doc.add_paragraph(f"Weighted Average Lease Term (WALT): {walt:.1f} years")
        doc.add_paragraph(f"Top Tenant Concentration: {top_pct*100:.1f}% of total revenue")

        if top5:
            doc.add_heading("Top Tenants by Revenue", level=2)
            tt_table = doc.add_table(rows=len(top5) + 1, cols=4)
            tt_table.style = 'Light Grid Accent 1'
            tt_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            tt_headers = ["Tenant", "SF", "Annual Revenue", "% of Total"]
            for i, h in enumerate(tt_headers):
                tt_table.rows[0].cells[i].text = h
                for paragraph in tt_table.rows[0].cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            total_rev = tenant_analytics.get("total_revenue", 0)
            for ri, t in enumerate(top5, 1):
                tt_table.rows[ri].cells[0].text = t.get("name", "")
                tt_table.rows[ri].cells[1].text = f"{t.get('_sf', 0):,.0f}"
                tt_table.rows[ri].cells[2].text = _fmt_currency(t.get("_annual_rev", 0))
                pct = t.get("_annual_rev", 0) / total_rev * 100 if total_rev > 0 else 0
                tt_table.rows[ri].cells[3].text = f"{pct:.1f}%"

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"Generated by CRE Lytic — {datetime.now().strftime('%Y-%m-%d %H:%M')}. "
                                "This memo is for informational purposes only.")
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = prop_name.replace(" ", "_").replace("/", "-")
    safe_name = safe_name.encode("ascii", "ignore").decode("ascii")[:40]
    filename = f"{safe_name}_Memo.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
